# pdf_to_docx_text.py
# -*- coding: utf-8 -*-

import io, os, sys
from tkinter import Tk, filedialog

import pdfplumber
from docx import Document
from docx.shared import Inches
from PIL import Image

# 設定
SLIDE_IMAGE_MAX_WIDTH_IN = 6.5  # 画像挿入時の最大幅（インチ）

def ask_pdf_path():
    root = Tk(); root.withdraw(); root.update()
    p = filedialog.askopenfilename(title="PDFファイルを選択してください", filetypes=[("PDF Files","*.pdf")])
    root.destroy()
    return p

def save_docx_from_pdf(pdf_path, docx_path):
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            # ページのテキストを抽出して段落に追加
            text = page.extract_text() or ""
            if text.strip():
                # ページ区切りの見出し（任意）
                doc.add_paragraph(f"--- Page {page_num} ---")
                for line in text.splitlines():
                    doc.add_paragraph(line)

            # 画像領域がある場合はページをレンダリングして切り出し、挿入
            if page.images:
                # page.to_image() を使って PIL Image を得る
                page_img = page.to_image(resolution=150)  # 解像度は必要に応じ調整
                pil = page_img.original.convert("RGB")
                for imginfo in page.images:
                    # pdfplumber の image dict: x0, top, x1, bottom (top=上からの距離)
                    x0 = imginfo.get("x0", 0)
                    top = imginfo.get("top", 0)
                    x1 = imginfo.get("x1", 0)
                    bottom = imginfo.get("bottom", 0)

                    # pt -> px (pdfplumber uses pt; to_image uses resolution)
                    pt_to_px = 150 / 72.0
                    left_px = int(x0 * pt_to_px)
                    top_px = int(top * pt_to_px)
                    right_px = int(x1 * pt_to_px)
                    bottom_px = int(bottom * pt_to_px)

                    # crop and insert
                    left_px = max(0, left_px); top_px = max(0, top_px)
                    right_px = min(pil.width, right_px); bottom_px = min(pil.height, bottom_px)
                    if right_px <= left_px or bottom_px <= top_px:
                        continue
                    crop = pil.crop((left_px, top_px, right_px, bottom_px))
                    bio = io.BytesIO()
                    crop.save(bio, format="PNG")
                    bio.seek(0)

                    # 画像幅をインチ換算してサイズ調整
                    img_width_in = (right_px - left_px) / 150.0  # px -> inch
                    if img_width_in > SLIDE_IMAGE_MAX_WIDTH_IN:
                        ratio = SLIDE_IMAGE_MAX_WIDTH_IN / img_width_in
                        # python-docx の add_picture で幅指定
                        doc.add_picture(bio, width=Inches(SLIDE_IMAGE_MAX_WIDTH_IN))
                    else:
                        doc.add_picture(bio)

    doc.save(docx_path)
    print("Saved:", docx_path)

def main():
    if len(sys.argv) >= 2:
        pdf_path = sys.argv[1]
    else:
        pdf_path = ask_pdf_path()
    if not pdf_path:
        print("PDF が選択されませんでした"); return
    if not os.path.exists(pdf_path):
        print("ファイルが見つかりません:", pdf_path); return
    docx_path = os.path.splitext(pdf_path)[0] + ".docx"
    save_docx_from_pdf(pdf_path, docx_path)

if __name__ == "__main__":
    main()
